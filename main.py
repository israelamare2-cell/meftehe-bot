import os
import telebot
from telebot import types
import google.generativeai as genai
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.section import WD_ORIENT
import io
import sqlite3
import threading
from flask import Flask
import time
import requests
import random  # አዲስ የተጨመረ - API Keys ለማፈራረቅ

# --- 1. ኮንፊገሬሽን ---
TELEGRAM_TOKEN = os.getenv('TELEGRAM_BOT_TOKEN')
CHANNEL_ID = "@digital_mat"

GITHUB_USER = "israelamare2-cell"
GITHUB_REPO = "meftehe-bot"
RELEASE_TAG = "v1" 

GITHUB_BASE_URL = f"https://github.com/{GITHUB_USER}/{GITHUB_REPO}/releases/download/{RELEASE_TAG}/"

# የ API Keys ዝግጅት (Load Balancing)
GEMINI_API_KEYS_STR = os.getenv('GEMINI_API_KEYS', os.getenv('GEMINI_API_KEY', ''))
API_KEY_LIST = [k.strip() for k in GEMINI_API_KEYS_STR.split(',')] if GEMINI_API_KEYS_STR else []

bot = telebot.TeleBot(TELEGRAM_TOKEN)
app = Flask(__name__)

# --- 2. ዳታቤዝ ---
def init_db():
    conn = sqlite3.connect('meftehe_national_data.db', check_same_thread=False)
    c = conn.cursor()
    c.execute('''CREATE TABLE IF NOT EXISTS exam_logs
                 (chat_id TEXT, subject TEXT, grade TEXT, exam_type TEXT,
                 sets TEXT, timestamp TEXT)''')
    conn.commit()
    conn.close()

init_db()

# --- 3. ዳታ ---
user_selection = {}
ALL_SUBJECTS = ["Mathematics", "Physics", "Chemistry", "Biology", "General Science", "English", "Social Studies", "Citizenship", "Amharic", "Afaan Oromoo", "Environmental Science", "Moral Education", "PVA", "HPE", "CTE", "Agriculture", "Economics", "IT"]
ALL_ASSESSMENT_TYPES = ["Mid Exam", "Final Exam", "Worksheet", "Quiz", "National Prep", "Model Exam", "Test"]

# አምስቱን ቋንቋዎች የያዘ የሜኑዎች ዲክሽነሪ
STRINGS = {
    'am': {
        'welcome': "እንኳን ወደ 'መፍትሔ' (Meftehe) በሰላም መጡ!",
        'start_btn': "🚀 ጀምር",
        'lang_msg': "🌍 ቋንቋ ይምረጡ / Select Language",
        'choose_mode': "🛠 **ምን ማዘጋጀት ይፈልጋሉ?**",
        'exam': "📝 የፈተና ዝግጅት",
        'note': "📚 የማስተማሪያ ኖት",
        'lesson': "📅 ዕለታዊ የትምህርት ዕቅድ",
        'review': "🔍 የመፅሀፍ ግምገማ",
        'back': "⬅️ ተመለስ",
        'sub_select': "📚 **ትምህርት ይምረጡ**",
        'gr_select': "📖 **የክፍል ደረጃ ይምረጡ**",
        'tp_select': "📝 **የፈተና አይነት ይምረጡ**",
        'diff_select': "📊 **የክብደት ደረጃ**",
        'ch_select': "📂 **ምዕራፍ ይምረጡ**",
        'set_select': "🛡️ **የሴት ብዛት**",
        'lang_option_msg': "📖 **ለቋንቋ ትምህርት ምን ይዘጋጅ?**",
        'opt_passage_only': "📄 ምንባብ ብቻ",
        'opt_question_only': "❓ ጥያቄ ብቻ",
        'opt_both': "📑 ምንባብ እና ጥያቄ"
    },
    'or': {
        'welcome': "Baga nagaan gara 'Meftehe' dhuftan!",
        'start_btn': "🚀 Jalqabi",
        'lang_msg': "🌍 Filannoo Afaanii",
        'choose_mode': "🛠 **Maal qopheessuu barbaaddu?**",
        'exam': "📝 Qophii Qorumsaa",
        'note': "📚 Nootii Barsiisaa",
        'lesson': "📅 Karoora Barnootaa",
        'review': "🔍 Gamaggama Kitaabaa",
        'back': "⬅️ Duubatti",
        'sub_select': "📚 **Gosa Barnootaa Filadhu**",
        'gr_select': "📖 **Kutaa Filadhu**",
        'tp_select': "📝 **Akaakuu Qorumsaa**",
        'diff_select': "📊 **Sadarkaa Ulfaatinaa**",
        'ch_select': "📂 **Boqonnaa Filadhu**",
        'set_select': "🛡️ **Baay'ina Seetii**",
        'lang_option_msg': "📖 **Barnoota afaaniif maal qopheessu?**",
        'opt_passage_only': "📄 Dubbisa Qofa",
        'opt_question_only': "❓ Gaaffii Qofa",
        'opt_both': "📑 Dubbisaafi Gaaffii"
    },
    'ti': {
        'welcome': "እንኳዕ ናብ 'መፍትሔ' (Meftehe) በሰላም መጻእኩም!",
        'start_btn': "🚀 ጀምር",
        'lang_msg': "🌍 ቋንቋ ይምረጡ",
        'choose_mode': "🛠 **እንታይ ክዳሎ ይደልዩ?**",
        'exam': "📝 ምድላው ፈተና",
        'note': "📚 ናይ መምህር ኖት",
        'lesson': "📅 ናይ መዓልቲ ትምህርቲ ትልሚ",
        'review': "🔍 ገምጋም መጽሓፍ",
        'back': "⬅️ ተመለስ",
        'sub_select': "📚 **ትምህርቲ ይምረጡ**",
        'gr_select': "📖 **ክፍሊ ይምረጡ**",
        'tp_select': "📝 **ዓይነት ፈተና ይምረጡ**",
        'diff_select': "📊 **ደረጃ ከቢድነት**",
        'ch_select': "📂 **ምዕራፍ ይምረጡ**",
        'set_select': "🛡️ **ብዝሒ ሴት**",
        'lang_option_msg': "📖 **ንትምህርቲ ቋንቋ እንታይ ይዳሎ?**",
        'opt_passage_only': "📄 ምንባብ ጥራሕ",
        'opt_question_only': "❓ ሕቶ ጥራሕ",
        'opt_both': "📑 ምንባብን ሕቶን"
    },
    'so': {
        'welcome': "Ku soo dhawaada 'Meftehe'!",
        'start_btn': "🚀 Bilow",
        'lang_msg': "🌍 Dooro Luqadda",
        'choose_mode': "🛠 **Maxaad rabtaa inaad diyaariso?**",
        'exam': "📝 Diyaarinta Imtixaanka",
        'note': "📚 Qoraalka Casharka",
        'lesson': "📅 Qorshe Cashar Maalinle",
        'review': "🔍 Dib-u-eegista Buugga",
        'back': "⬅️ Dib u laabo",
        'sub_select': "📚 **Dooro Maaddada**",
        'gr_select': "📖 **Dooro Fasalka**",
        'tp_select': "📝 **Dooro Nooca Imtixaanka**",
        'diff_select': "📊 **Heerka Adkaanta**",
        'ch_select': "📂 **Dooro Cutubka**",
        'set_select': "🛡️ **Tirada Isu-geynta**",
        'lang_option_msg': "📖 **Maxaa loo diyaariyaa casharka luqadda?**",
        'opt_passage_only': "📄 Qoraalka kaliya",
        'opt_question_only': "❓ Su'aalaha kaliya",
        'opt_both': "📑 Qoraalka iyo Su'aalaha"
    },
    'en': {
        'welcome': "Welcome to 'Meftehe'!",
        'start_btn': "🚀 Start",
        'lang_msg': "🌍 Select Language",
        'choose_mode': "🛠 **What would you like to prepare?**",
        'exam': "📝 Exam Preparation",
        'note': "📚 Teacher Notes",
        'lesson': "📅 Daily Lesson Plan",
        'review': "🔍 Textbook Review",
        'back': "⬅️ Back",
        'sub_select': "📚 **Select Subject**",
        'gr_select': "📖 **Select Grade**",
        'tp_select': "📝 **Select Exam Type**",
        'diff_select': "📊 **Difficulty Level**",
        'ch_select': "📂 **Select Chapter**",
        'set_select': "🛡️ **Number of Sets**",
        'lang_option_msg': "📖 **What to prepare for Language subject?**",
        'opt_passage_only': "📄 Passage Only",
        'opt_question_only': "❓ Questions Only",
        'opt_both': "📑 Passage & Questions"
    }
}

# --- 4. ረዳት ፈንክሽኖች ---
def is_subscribed(user_id):
    try:
        member = bot.get_chat_member(CHANNEL_ID, user_id)
        return member.status in ['member', 'administrator', 'creator']
    except:
        return False

def download_book_from_github(grade, subject):
    if not os.path.exists("books"):
        os.makedirs("books")
    
    subject_filename = subject.lower().replace(" ", "_")
    filename = f"grade{grade}_{subject_filename}.pdf"
    local_path = f"books/{filename}"
    
    if os.path.exists(local_path):
        return local_path
    
    url = f"{GITHUB_BASE_URL}{filename}"
    try:
        response = requests.get(url, stream=True, timeout=20)
        if response.status_code == 200:
            with open(local_path, 'wb') as f:
                for chunk in response.iter_content(chunk_size=8192):
                    f.write(chunk)
            return local_path
        return None
    except Exception as e:
        print(f"Download error: {e}")
        return None

# --- 5. የቦት ሜኑዎች ---

@bot.message_handler(commands=['start'])
def start(message):
    user_id = message.from_user.id
    chat_id = message.chat.id
    
    if not is_subscribed(user_id):
        markup = types.InlineKeyboardMarkup()
        markup.add(types.InlineKeyboardButton("📢 ቻናሉን ተቀላቀል (Join)", url="https://t.me/digital_mat"))
        markup.add(types.InlineKeyboardButton("✅ Verify", callback_data="check_subs"))
        bot.send_message(chat_id, "⚠️ ቦቱን ለመጠቀም መጀመሪያ ቻናላችንን ይቀላቀሉ!", reply_markup=markup)
        return

    user_selection[chat_id] = {'counts': {}, 'lang': 'am'}
    markup = types.InlineKeyboardMarkup(row_width=2)
    btns = [
        types.InlineKeyboardButton("አማርኛ", callback_data="lang_am"),
        types.InlineKeyboardButton("Afaan Oromoo", callback_data="lang_or"),
        types.InlineKeyboardButton("ትግርኛ", callback_data="lang_ti"),
        types.InlineKeyboardButton("Af-Soomaali", callback_data="lang_so"),
        types.InlineKeyboardButton("English", callback_data="lang_en")
    ]
    markup.add(*btns)
    bot.send_message(chat_id, "🌍 **Select Language / ቋንቋ ይምረጡ**", reply_markup=markup, parse_mode="Markdown")

@bot.callback_query_handler(func=lambda call: True)
def handle_callbacks(call):
    chat_id = call.message.chat.id
    data = call.data
    
    if data.startswith("lang_"):
        lang_code = data.split('_')[1]
        user_selection[chat_id]['lang'] = lang_code
        lang = lang_code
        
        welcome_text = (
            f"🌟 **{STRINGS[lang]['welcome']}** 🌟\n\n"
            "መምህራንን በዘመናዊ AI በማገዝ የትምህርት ጥራትን ለማረጋገጥ የተፈጠረ ሁለገብ የዲጂታል ረዳት።\n\n"
            f"🛠 **{STRINGS[lang]['choose_mode']}**"
        )
        
        markup = types.InlineKeyboardMarkup()
        markup.add(types.InlineKeyboardButton(STRINGS[lang]['start_btn'], callback_data="mode_selection"))
        bot.edit_message_text(welcome_text, chat_id, call.message.message_id, reply_markup=markup, parse_mode="Markdown")
        return

    lang = user_selection.get(chat_id, {}).get('lang', 'am')

    if data == "check_subs":
        if is_subscribed(call.from_user.id): start(call.message)
        else: bot.answer_callback_query(call.id, "❌ አባል አይደሉም!", show_alert=True)

    elif data == "mode_selection":
        markup = types.InlineKeyboardMarkup(row_width=2)
        markup.add(types.InlineKeyboardButton(STRINGS[lang]['exam'], callback_data="set_mode_exam"),
                   types.InlineKeyboardButton(STRINGS[lang]['note'], callback_data="set_mode_note"),
                   types.InlineKeyboardButton(STRINGS[lang]['lesson'], callback_data="set_mode_lesson"),
                   types.InlineKeyboardButton(STRINGS[lang]['review'], callback_data="set_mode_review"))
        markup.add(types.InlineKeyboardButton(STRINGS[lang]['back'], callback_data="start_back"))
        bot.edit_message_text(STRINGS[lang]['choose_mode'], chat_id, call.message.message_id, reply_markup=markup, parse_mode="Markdown")

    elif data == "start_back":
        start(call.message)

    elif data.startswith("set_mode_"):
        user_selection[chat_id]['mode'] = data.split('_')[2]
        markup = types.InlineKeyboardMarkup(row_width=2)
        btns = [types.InlineKeyboardButton(s, callback_data=f"sub_{s}") for s in ALL_SUBJECTS]
        markup.add(*btns)
        markup.add(types.InlineKeyboardButton(STRINGS[lang]['back'], callback_data="mode_selection"))
        bot.edit_message_text(STRINGS[lang]['sub_select'], chat_id, call.message.message_id, reply_markup=markup, parse_mode="Markdown")

    elif data == "main_menu":
        call.data = "mode_selection"
        handle_callbacks(call)

    elif data.startswith('sub_'):
        user_selection[chat_id]['subject'] = data.split('_')[1]
        markup = types.InlineKeyboardMarkup(row_width=4)
        btns = [types.InlineKeyboardButton(f"{i}ኛ", callback_data=f"gr_{i}") for i in range(1, 13)]
        markup.add(*btns, types.InlineKeyboardButton(STRINGS[lang]['back'], callback_data="main_menu"))
        bot.edit_message_text(STRINGS[lang]['gr_select'], chat_id, call.message.message_id, reply_markup=markup, parse_mode="Markdown")

    elif data.startswith('gr_'):
        user_selection[chat_id]['grade'] = data.split('_')[1]
        mode = user_selection[chat_id]['mode']
        if mode == "exam":
            markup = types.InlineKeyboardMarkup(row_width=2)
            btns = [types.InlineKeyboardButton(t, callback_data=f"tp_{t}") for t in ALL_ASSESSMENT_TYPES]
            markup.add(*btns, types.InlineKeyboardButton(STRINGS[lang]['back'], callback_data=f"sub_{user_selection[chat_id]['subject']}"))
            bot.edit_message_text(STRINGS[lang]['tp_select'], chat_id, call.message.message_id, reply_markup=markup, parse_mode="Markdown")
        elif mode == "review":
            call.data = "tp_Review"
            handle_callbacks(call)
        elif mode == "lesson":
            call.data = "tp_LessonPlan"
            handle_callbacks(call)
        else:
            call.data = "tp_Note"
            handle_callbacks(call)

    elif data.startswith('tp_'):
        user_selection[chat_id]['type'] = data.split('_')[1]
        mode = user_selection[chat_id]['mode']
        if mode == "exam":
            markup = types.InlineKeyboardMarkup(row_width=2)
            markup.add(types.InlineKeyboardButton("ቀላል", callback_data="df_Easy"), types.InlineKeyboardButton("መካከለኛ", callback_data="df_Medium"),
                       types.InlineKeyboardButton("ከባድ", callback_data="df_Hard"), types.InlineKeyboardButton("⚖️ Mixed", callback_data="df_MixedFair"))
            markup.add(types.InlineKeyboardButton(STRINGS[lang]['back'], callback_data=f"gr_{user_selection[chat_id]['grade']}"))
            bot.edit_message_text(STRINGS[lang]['diff_select'], chat_id, call.message.message_id, reply_markup=markup, parse_mode="Markdown")
        elif mode == "review":
            call.data = "df_Standard"
            handle_callbacks(call)
        elif mode == "lesson":
            call.data = "df_SMASE"
            handle_callbacks(call)
        else:
            call.data = "df_Normal"
            handle_callbacks(call)

    elif data.startswith('df_'):
        user_selection[chat_id]['diff'] = data.split('_')[1]
        mode = user_selection[chat_id]['mode']
        if mode == "exam":
            levels = ["Knowledge", "Understanding", "Application", "Analysis", "Evaluation", "Creation", "Mixed"]
            markup = types.InlineKeyboardMarkup(row_width=2)
            btns = [types.InlineKeyboardButton(l, callback_data=f"bl_{l}") for l in levels]
            markup.add(*btns, types.InlineKeyboardButton(STRINGS[lang]['back'], callback_data=f"tp_{user_selection[chat_id]['type']}"))
            bot.edit_message_text("🧠 **Bloom's Taxonomy**", chat_id, call.message.message_id, reply_markup=markup)
        elif mode == "review":
            call.data = "bl_Audit"
            handle_callbacks(call)
        elif mode == "lesson":
            call.data = "bl_ActiveLearning"
            handle_callbacks(call)
        else:
            call.data = "bl_Note"
            handle_callbacks(call)

    elif data.startswith('bl_'):
        user_selection[chat_id]['bloom'] = data.split('_')[1]
        markup = types.InlineKeyboardMarkup(row_width=3)
        btns = [types.InlineKeyboardButton(f"ምዕ {i}", callback_data=f"ch_{i}") for i in range(1, 10)]
        markup.add(*btns)
        markup.add(types.InlineKeyboardButton("📚 All", callback_data="ch_all"), 
                   types.InlineKeyboardButton("🤖 በራሱ (Auto)", callback_data="ch_auto"),
                   types.InlineKeyboardButton("✏️ በእጅ ጻፍ", callback_data="manual_chapter"))
        
        back_data = f"df_{user_selection[chat_id]['diff']}" if user_selection[chat_id]['mode'] == "exam" else f"sub_{user_selection[chat_id]['subject']}"
        markup.add(types.InlineKeyboardButton(STRINGS[lang]['back'], callback_data=back_data))
        bot.edit_message_text(STRINGS[lang]['ch_select'], chat_id, call.message.message_id, reply_markup=markup, parse_mode="Markdown")

    elif data == "manual_chapter":
        msg = bot.send_message(chat_id, "🖋 **እባክዎ የሚፈልጉትን ምዕራፍ ወይም ርዕስ ይጻፉ፡**")
        bot.register_next_step_handler(msg, process_manual_chapter)

    elif data.startswith('ch_'):
        user_selection[chat_id]['chapter'] = data.split('_')[1]
        mode = user_selection[chat_id]['mode']
        subject = user_selection[chat_id]['subject']
        is_language = subject.lower() in ["amharic", "english", "afaan oromoo"]
        
        if mode == "lesson":
            markup = types.InlineKeyboardMarkup(row_width=2)
            markup.add(
                types.InlineKeyboardButton("📄 አንድ ገፅ ብቻ", callback_data="pgtype_single"),
                types.InlineKeyboardButton("📖 የገፅ ክልል (ከ-እስ)", callback_data="pgtype_range")
            )
            bot.edit_message_text("🔢 **ገፅ እንዴት መምረጥ ይፈልጋሉ?**", chat_id, call.message.message_id, reply_markup=markup)

        elif mode == "exam":
            if is_language:
                markup = types.InlineKeyboardMarkup(row_width=1)
                markup.add(
                    types.InlineKeyboardButton(STRINGS[lang]['opt_passage_only'], callback_data="lopt_PassageOnly"),
                    types.InlineKeyboardButton(STRINGS[lang]['opt_question_only'], callback_data="lopt_QuestionOnly"),
                    types.InlineKeyboardButton(STRINGS[lang]['opt_both'], callback_data="lopt_Both")
                )
                bot.edit_message_text(STRINGS[lang]['lang_option_msg'], chat_id, call.message.message_id, reply_markup=markup)
            else:
                markup = types.InlineKeyboardMarkup()
                markup.add(types.InlineKeyboardButton("🔓 1 Set", callback_data="sec_1"), 
                           types.InlineKeyboardButton("🛡️ 2 Sets", callback_data="sec_2"),
                           types.InlineKeyboardButton("🛡️ 4 Sets", callback_data="sec_4"))
                markup.add(types.InlineKeyboardButton(STRINGS[lang]['back'], callback_data=f"bl_{user_selection[chat_id]['bloom']}"))
                bot.edit_message_text(STRINGS[lang]['set_select'], chat_id, call.message.message_id, reply_markup=markup, parse_mode="Markdown")
        elif mode == "review":
            markup = types.InlineKeyboardMarkup(row_width=1)
            markup.add(types.InlineKeyboardButton("🔍 አጠቃላይ ኦዲት (Full Audit)", callback_data="rev_FullAudit"),
                       types.InlineKeyboardButton("🧠 ፔዳጎጂ እና SMASE ግምገማ", callback_data="rev_Pedagogy"),
                       types.InlineKeyboardButton("📝 የጥያቄዎች ጥራት ኦዲት", callback_data="rev_Assessment"),
                       types.InlineKeyboardButton("🇪🇹 የሀገር በቀል እውቀትና እሴት", callback_data="rev_Indigenous"),
                       types.InlineKeyboardButton("🚀 የ21ኛው ክፍለዘመን ክህሎቶች", callback_data="rev_21stCentury"),
                       types.InlineKeyboardButton("🌍 አካታችነትና የፆታ ተመጣጣኝነት", callback_data="rev_Inclusivity"))
            markup.add(types.InlineKeyboardButton(STRINGS[lang]['back'], callback_data=f"bl_{user_selection[chat_id]['bloom']}"))
            bot.edit_message_text("🧐 **የግምገማ ዘርፍ ይምረጡ**", chat_id, call.message.message_id, reply_markup=markup)
        else:
            markup = types.InlineKeyboardMarkup(row_width=2)
            markup.add(types.InlineKeyboardButton("1. 🎯 አላማና መግቢያ", callback_data="nt_1_Objectives"),
                       types.InlineKeyboardButton("2. 📖 ዝርዝር ማስታወሻ", callback_data="nt_2_Comprehensive"),
                       types.InlineKeyboardButton("3. 💡 ማብራሪያና ምሳሌ", callback_data="nt_3_Examples"),
                       types.InlineKeyboardButton("4. 📝 አጭር ማጠቃለያ", callback_data="nt_4_Summary"),
                       types.InlineKeyboardButton("5. 🧠 የክለሳ ጥያቄዎች", callback_data="nt_5_ReviewQs"),
                       types.InlineKeyboardButton("6. 🚀 ሁሉንም በአንድ", callback_data="nt_6_FullPackage"))
            markup.add(types.InlineKeyboardButton("🔀 ምርጫህን አሰባጥር (Custom Mix)", callback_data="nt_custom_mix"))
            markup.add(types.InlineKeyboardButton(STRINGS[lang]['back'], callback_data=f"bl_{user_selection[chat_id]['bloom']}"))
            bot.edit_message_text("✨ **የኖት አይነት ይምረጡ ወይም ስብጥር ያዝዙ**", chat_id, call.message.message_id, reply_markup=markup)

    elif data.startswith('pgtype_'):
        user_selection[chat_id]['page_type'] = data.split('_')[1]
        msg_text = "🖋 **እባክዎ የገፅ ቁጥሩን ይጻፉ (ለምሳሌ፦ 12):**" if data == "pgtype_single" else "🖋 **እባክዎ የገፅ ክልሉን ይጻፉ (ለምሳሌ፦ 12-15):**"
        msg = bot.send_message(chat_id, msg_text)
        bot.register_next_step_handler(msg, final_generation_trigger)

    elif data.startswith('lopt_'):
        user_selection[chat_id]['lang_output_option'] = data.split('_')[1]
        markup = types.InlineKeyboardMarkup()
        markup.add(types.InlineKeyboardButton("🔓 1 Set", callback_data="sec_1"), 
                   types.InlineKeyboardButton("🛡️ 2 Sets", callback_data="sec_2"),
                   types.InlineKeyboardButton("🛡️ 4 Sets", callback_data="sec_4"))
        bot.edit_message_text(STRINGS[lang]['set_select'], chat_id, call.message.message_id, reply_markup=markup)

    elif data.startswith('sec_'):
        user_selection[chat_id]['num_sets'] = int(data.split('_')[1])
        msg = bot.send_message(chat_id, "🔢 **የጥያቄ ብዛትና መዋቅር ይጻፉ**\n(ለምሳሌ፦ 'ምርጫ=10, እውነት/ሐሰት=5')\nወይም 'auto' ብለው ይጻፉ።")
        bot.register_next_step_handler(msg, final_generation_trigger)

    elif data.startswith('rev_'):
        user_selection[chat_id]['review_type'] = data.split('_')[1]
        markup = types.InlineKeyboardMarkup(row_width=2)
        ranges = ["1-20", "21-50", "51-100", "101-150", "151-200", "All Pages"]
        btns = [types.InlineKeyboardButton(f"ገፅ {r}", callback_data=f"pg_{r}") for r in ranges]
        markup.add(*btns)
        markup.add(types.InlineKeyboardButton(STRINGS[lang]['back'], callback_data=f"ch_{user_selection[chat_id]['chapter']}"))
        bot.edit_message_text("📄 **ለመገምገም የሚፈልጉትን የገፅ ክልል ይምረጡ፦**", chat_id, call.message.message_id, reply_markup=markup)

    elif data.startswith('pg_'):
        user_selection[chat_id]['page_range'] = data.split('_')[1]
        msg = bot.send_message(chat_id, "🖋 **ልዩ ትኩረት እንዲሰጥበት የሚፈልጉት ነጥብ ካለ ይጻፉ (ወይም 'auto')፡**")
        bot.register_next_step_handler(msg, final_generation_trigger)

    elif data == "nt_custom_mix":
        msg = bot.send_message(chat_id, "🔢 **የሚፈልጓቸውን የኖት አይነቶች ቁጥሮች በኮማ ለይተው ይጻፉ፡**\nለምሳሌ፦ '1, 3, 5'")
        bot.register_next_step_handler(msg, process_custom_mix)

    elif data.startswith('nt_'):
        user_selection[chat_id]['note_style'] = data.split('_')[2] 
        style = user_selection[chat_id]['note_style']
        if "ReviewQs" in style or "FullPackage" in style:
            markup = types.InlineKeyboardMarkup(row_width=2)
            btns = [types.InlineKeyboardButton(t, callback_data=f"nrtp_{t}") for t in ALL_ASSESSMENT_TYPES]
            markup.add(*btns, types.InlineKeyboardButton(STRINGS[lang]['back'], callback_data=f"ch_{user_selection[chat_id]['chapter']}"))
            bot.edit_message_text("📝 **የክለሳ ጥያቄዎቹ ምን አይነት ይሁኑ?**", chat_id, call.message.message_id, reply_markup=markup)
        else:
            msg = bot.send_message(chat_id, "📄 **ማስታወሻው እንዲያካትት የሚፈልጉት ልዩ ነጥብ ካለ ይጻፉ (ወይም 'auto')፡**")
            bot.register_next_step_handler(msg, final_generation_trigger)

    elif data.startswith('nrtp_'):
        user_selection[chat_id]['nr_type'] = data.split('_')[1]
        markup = types.InlineKeyboardMarkup(row_width=2)
        markup.add(types.InlineKeyboardButton("ቀላል", callback_data="nrdf_Easy"), types.InlineKeyboardButton("መካከለኛ", callback_data="nrdf_Medium"),
                   types.InlineKeyboardButton("ከባድ", callback_data="nrdf_Hard"), types.InlineKeyboardButton("⚖️ Mixed", callback_data="nrdf_MixedFair"))
        bot.edit_message_text("📊 **የክለሳ ጥያቄዎቹ የክብደት ደረጃ**", chat_id, call.message.message_id, reply_markup=markup)

    elif data.startswith('nrdf_'):
        user_selection[chat_id]['nr_diff'] = data.split('_')[1]
        levels = ["Knowledge", "Understanding", "Application", "Analysis", "Evaluation", "Creation", "Mixed"]
        markup = types.InlineKeyboardMarkup(row_width=2)
        btns = [types.InlineKeyboardButton(l, callback_data=f"nrbl_{l}") for l in levels]
        markup.add(*btns)
        bot.edit_message_text("🧠 **የጥያቄዎቹ Bloom's Taxonomy**", chat_id, call.message.message_id, reply_markup=markup)

    elif data.startswith('nrbl_'):
        user_selection[chat_id]['nr_bloom'] = data.split('_')[1]
        markup = types.InlineKeyboardMarkup()
        markup.add(types.InlineKeyboardButton("🔓 1 Set", callback_data="nrsec_1"), types.InlineKeyboardButton("🛡️ 2 Sets", callback_data="nrsec_2"),
                   types.InlineKeyboardButton("🛡️ 4 Sets", callback_data="nrsec_4"))
        bot.edit_message_text("🛡️ **የክለሳ ጥያቄዎቹ የሴት ብዛት**", chat_id, call.message.message_id, reply_markup=markup)

    elif data.startswith('nrsec_'):
        user_selection[chat_id]['nr_sets'] = int(data.split('_')[1])
        msg = bot.send_message(chat_id, "🔢 **የጥያቄ ብዛትና መዋቅር ይጻፉ (ወይም 'auto')፡**")
        bot.register_next_step_handler(msg, final_generation_trigger)

# --- 6. AI Generation Helpers ---

def process_custom_mix(message):
    chat_id = message.chat.id
    user_selection[chat_id]['note_style'] = f"Custom Mix: {message.text}"
    if "5" in message.text or "6" in message.text:
        markup = types.InlineKeyboardMarkup(row_width=2)
        btns = [types.InlineKeyboardButton(t, callback_data=f"nrtp_{t}") for t in ALL_ASSESSMENT_TYPES]
        markup.add(*btns)
        bot.send_message(chat_id, "📝 **በስብጥሩ ውስጥ ለሚገኙት የክለሳ ጥያቄዎች ምን አይነት ይሁኑ?**", reply_markup=markup)
    else:
        msg = bot.send_message(chat_id, "📄 **ልዩ ነጥብ ካለ ይጻፉ (ወይም 'auto')፡**")
        bot.register_next_step_handler(msg, final_generation_trigger)

def process_manual_chapter(message):
    chat_id = message.chat.id
    user_selection[chat_id]['chapter'] = message.text
    mode = user_selection[chat_id]['mode']
    lang = user_selection[chat_id].get('lang', 'am')
    if mode == "exam":
        markup = types.InlineKeyboardMarkup()
        markup.add(types.InlineKeyboardButton("🔓 1 Set", callback_data="sec_1"), types.InlineKeyboardButton("🛡️ 2 Sets", callback_data="sec_2"),
                   types.InlineKeyboardButton("🛡️ 4 Sets", callback_data="sec_4"))
        bot.send_message(chat_id, STRINGS[lang]['set_select'], reply_markup=markup)
    elif mode == "review":
        markup = types.InlineKeyboardMarkup(row_width=1)
        markup.add(types.InlineKeyboardButton("🔍 አጠቃላይ ኦዲት (Full Audit)", callback_data="rev_FullAudit"),
                   types.InlineKeyboardButton("🧠 ፔዳጎጂ እና SMASE ግምገማ", callback_data="rev_Pedagogy"),
                   types.InlineKeyboardButton("📝 የጥያቄዎች ጥራት ኦዲት", callback_data="rev_Assessment"))
        bot.send_message(chat_id, "🧐 **የግምገማ ዘርፍ ይምረጡ**", reply_markup=markup)
    elif mode == "lesson":
        markup = types.InlineKeyboardMarkup(row_width=2)
        markup.add(
            types.InlineKeyboardButton("📄 አንድ ገፅ ብቻ", callback_data="pgtype_single"),
            types.InlineKeyboardButton("📖 የገፅ ክልል (ከ-እስ)", callback_data="pgtype_range")
        )
        bot.send_message(chat_id, "🔢 **ገፅ እንዴት መምረጥ ይፈልጋሉ?**", reply_markup=markup)
    else:
        markup = types.InlineKeyboardMarkup(row_width=2)
        markup.add(types.InlineKeyboardButton("1. 🎯 አላማና መግቢያ", callback_data="nt_1_Objectives"),
                   types.InlineKeyboardButton("2. 📖 ዝርዝር ማስታወሻ", callback_data="nt_2_Comprehensive"),
                   types.InlineKeyboardButton("3. 💡 ማብራሪያና ምሳሌ", callback_data="nt_3_Examples"),
                   types.InlineKeyboardButton("4. 📝 አጭር ማጠቃለያ", callback_data="nt_4_Summary"),
                   types.InlineKeyboardButton("5. 🧠 የክለሳ ጥያቄዎች", callback_data="nt_5_ReviewQs"),
                   types.InlineKeyboardButton("6. 🚀 ሁሉንም በአንድ", callback_data="nt_6_FullPackage"))
        bot.send_message(chat_id, "✨ **የኖት አይነት ይምረጡ**", reply_markup=markup)

def final_generation_trigger(message):
    user_selection[message.chat.id]['tos_config'] = message.text
    generate_final_content(message)

def generate_final_content(message):
    chat_id = message.chat.id
    data = user_selection.get(chat_id)
    if not data: return

    bot.send_message(chat_id, "🔍 መፅሀፉን ከ GitHub ማከማቻዬ እየፈለግኩ ነው...")
    file_path = download_book_from_github(data['grade'], data['subject'])
    
    if not file_path:
        bot.send_message(chat_id, "❌ መፅሀፉ በ GitHub Release ላይ አልተገኘም!")
        return

    bot.send_message(chat_id, f"🚀 መጽሐፉን አግኝቻለሁ! አሁን {data['mode']} እያዘጋጀሁ ነው...")
    
    try:
        selected_lang = data.get('lang', 'am')
        lang_map = {
            'am': "STRICTLY in AMHARIC language.",
            'or': "STRICTLY in AFAAN OROMOO language using professional terms.",
            'ti': "STRICTLY in TIGRIGNA language.",
            'so': "STRICTLY in SOMALI language.",
            'en': "STRICTLY in ENGLISH language."
        }
        
        target_subject = data['subject'].lower()
        if target_subject == "afaan oromoo":
            lang_rule = lang_map['or']
        elif target_subject == "amharic":
            lang_rule = lang_map['am']
        elif target_subject == "english":
            lang_rule = lang_map['en']
        else:
            lang_rule = lang_map[selected_lang]

        lang_output_instruction = ""
        if 'lang_output_option' in data:
            opt = data['lang_output_option']
            if opt == "PassageOnly":
                lang_output_instruction = "IMPORTANT: Provide ONLY the Reading Passage based on the context. Do NOT generate any questions."
            elif opt == "QuestionOnly":
                lang_output_instruction = "IMPORTANT: Generate ONLY the questions. Do NOT include the reading passage text itself."
            elif opt == "Both":
                lang_output_instruction = "IMPORTANT: Provide the Reading Passage FIRST, followed by the related questions."

        if data['mode'] == "lesson":
            prompt = f"""You are a Professional Curriculum Developer specializing in SMASE (Active Learning).
            TASK: Create a DAILY LESSON PLAN based on Chapter: {data['chapter']} and Page: {data.get('tos_config', 'auto')}.
            
            STRICT REQUIREMENTS:
            1. LANGUAGE: {lang_rule}
            2. PEDAGOGY: Follow SMASE (Active Learning). Ensure it's learner-centered.
            3. STYLE: Use VERY SHORT BULLET POINTS. NO long sentences.
            4. FORMAT: Use a CLEAR TABLE for the Teacher/Student activity sections.
            
            HEADER INFO:
            - School: የካ ተራራ ቅድመ አንደኛ፣ አንደኛ እና መካከለኛ ደረጃ ትምህርት ቤት
            - Teacher: እስራኤል አማረ
            
            SECTIONS TO INCLUDE (Short & Precise):
            - Objectives (አላማዎች)
            - Significance (አስፈላጊነት)
            - Prior Knowledge (ቀዳሚ ዕዉቀት)
            - Competency (አጥጋቢ የመማር ብቃት)
            
            TABLE STRUCTURE:
            Generate a table with these columns: [የመማር ማስተማር ቅደም ተከተል, ክፍለ ጊዜ, ይዘት, የመምህሩ ተግባር, የተማሪ ተግባር, ምዘና, የመርጃ መሣሪያ]
            (Fill the columns with short bulleted points like: • ሰላምታ • ክለሳ • ገለጻ)
            
            DIFFERENTIATED SUPPORT (Short examples):
            - For High Achievers (ላቅ ባለ ደረጃ ላሉ)
            - For Average Students (በመካከለኛ ደረጃ ላሉ)
            - For Low Achievers (ዝቅ ባለ ደረጃ ላሉ)
            - Special Needs (ልዩ ፍላጎት)
            """

        elif data['mode'] == "exam":
            prompt = f"""You are an expert Ethiopian National Examiner.
            STRICT COMPLIANCE:
            1. SOURCE: Use ONLY the provided PDF. Focus on Chapter: {data['chapter']}, Bloom Level: {data['bloom']}, Difficulty: {data['diff']}.
            2. USER COMMAND: Create exam structure: {data['tos_config']}. If 'auto', decide a standard structure.
            3. LANGUAGE: {lang_rule}
            4. SYMBOLS: ALL formulas in LaTeX using $inline$ or $$display$$.
            5. LANGUAGE SUBJECT RULE: {lang_output_instruction}
            6. OUTPUT: {data['num_sets']} different sets. Include TOS, Exam, and Answer Key. Page break using '---PAGE BREAK---'."""
        
        elif data['mode'] == "review":
            review_type = data['review_type']
            page_range = data.get('page_range', 'specified chapters')
            prompt = f"""You are a Precise Curriculum Auditor.
            TASK: Conduct a PAGE-BY-PAGE Audit of the PDF for Page Range/Chapter: {page_range} / {data['chapter']}.
            REVIEW SCOPE: {review_type}
            
            STRICT OUTPUT STRUCTURE:
            1. EXECUTIVE SUMMARY: A brief overview of the quality.
            2. DETAILED PAGE-BY-PAGE FINDINGS: 
               - Format: [Page X]: List specific errors (factual, grammatical, or pedagogical) or improvement points.
            3. CRITICAL ERRORS TABLE: A table showing [Page #], [Current Content], [Suggested Correction].
            4. PEDAGOGICAL ALIGNMENT: How it fits SMASE and 21st Century Skills.
            
            FORMATTING:
            - Use Bold for Page numbers.
            - Use Tables for corrections using '|' symbols.
            - LANGUAGE: {lang_rule}
            - USER SPECIAL NOTE: {data.get('tos_config', 'auto')}"""

        else:
            style_request = data.get('note_style', data.get('tos_config', 'FullPackage'))
            prompt = f"Professional Curriculum Expert Note Generation for Chapter {data['chapter']}... Style: {style_request}. Language: {lang_rule}..."

        with open(file_path, "rb") as f:
            file_data = f.read()

        # ==============================================================
        # ANTI-CRASH (RATE LIMIT) ማስተካከያ የተጨመረበት ክፍል
        # ==============================================================
        max_retries = 4
        response = None
        
        for attempt in range(max_retries):
            try:
                # 1. በዘፈቀደ ቁልፍ መምረጥ (Load Balancing)
                if API_KEY_LIST:
                    current_key = random.choice(API_KEY_LIST)
                    genai.configure(api_key=current_key)
                
                model = genai.GenerativeModel('gemini-2.5-flash')
                
                # 2. AIውን ማዘዝ
                response = model.generate_content([{"mime_type": "application/pdf", "data": file_data}, prompt])
                break # ተሳክቷል! ከ Loop ውጣ
                
            except Exception as e:
                error_text = str(e).lower()
                # ኤረሩ የትራፊክ መጨናነቅ ከሆነ
                if "429" in error_text or "quota" in error_text or "rate limit" in error_text or "503" in error_text:
                    if attempt < max_retries - 1:
                        wait_time = 5 * (attempt + 1)
                        bot.send_message(chat_id, f"⏳ የሰዎች መብዛት (Traffic) አጋጥሟል። ከ {wait_time} ሰከንድ በኋላ በድጋሚ እየሞከርኩ ነው፣ እባክዎ ይጠብቁ...")
                        time.sleep(wait_time)
                        continue
                # የትራፊክ ካልሆነ ወይም ሙከራው ካለቀ ኤረሩን አውጣ
                raise Exception(f"ከ{max_retries} ሙከራዎች በኋላ አልተሳካም። እባክዎ ትንሽ ቆይተው ይሞክሩ። (Error: {str(e)[:50]})")
        # ==============================================================

        raw_content = response.text.replace("###", "").replace("##", "")
        
        doc = Document()
        
        if data['mode'] == "lesson":
            # ዕለታዊ የትምህርት ዕቅድ ፎርማት
            section = doc.sections[0]
            section.orientation = WD_ORIENT.LANDSCAPE
            new_width, new_height = section.page_height, section.page_width
            section.page_width = new_width
            section.page_height = new_height

            header = doc.add_paragraph("ዕለታዊ የትምህርት ዕቅድ", style='Header')
            header.alignment = 1
            
            info = doc.add_paragraph()
            info.add_run(f"የመምህሩ ስም: እስራኤል አማረ\t\t\t\tየት/ቤቱ ስም: የካ ተራራ ቅድመ አንደኛ፣ አንደኛ እና መካከለኛ ደረጃ ትምህርት ቤት\n")
            info.add_run(f"የትም ዓይነት: {data['subject']}\t\t\t\tምዕራፍ: {data['chapter']}\n")
            info.add_run(f"የክፍል ደረጃ: {data['grade']}\t\t\t\tየዕለቱ ገፅ: {data.get('tos_config', '')}")

            # AI የመለሰውን ይዘት (ሰንጠረዡን ጨምሮ) ወደ ሰነዱ ማዛወር
            doc.add_paragraph("\n" + raw_content)
            
            # ፊርማ
            doc.add_paragraph("\nመምህር: እስራኤል አማረ _________ \t የት/ክፍል ተጠሪ: አስመራወርቅ ሀይሌ _________ \t ም/ር/መ/ር: ከበደ ተስፋዪ _________")
        
        else:
            title = doc.add_heading(f"{data['subject']} - Grade {data['grade']} {data['mode'].upper()}", 0)
            title.alignment = 1 

            if data['mode'] == "review":
                meta = doc.add_paragraph()
                meta.add_run(f"Audit Scope: {data['review_type']}\nPages Reviewed: {data.get('page_range', 'All')}\n").bold = True

            sections = raw_content.split('\n\n')
            for section in sections:
                clean_sec = section.strip()
                if not clean_sec: continue
                
                if clean_sec.startswith("[Page") or clean_sec.startswith("Page") or ":" in clean_sec.split('\n')[0]:
                    p = doc.add_paragraph()
                    run = p.add_run(clean_sec)
                    run.bold = True
                elif "|" in clean_sec: 
                    doc.add_paragraph(clean_sec) 
                else:
                    doc.add_paragraph(clean_sec)
        
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        file_stream.name = f"{data['subject']}_{data['mode']}.docx"
        bot.send_document(chat_id, file_stream, caption=f"✅ {data['mode'].capitalize()}ው በተሳካ ሁኔታ ተዘጋጅቷል።")

    except Exception as e:
        bot.send_message(chat_id, f"❌ ስህተት ተፈጥሯል፦ {str(e)}")

# --- 7. ሰርቨር እና ቦት ማስነሻ ---
@app.route('/')
def home(): 
    return "Meftehe Bot is Online!"

def run_bot():
    while True:
        try:
            bot.remove_webhook()
            time.sleep(2)
            bot.infinity_polling(skip_pending=True, timeout=60)
        except Exception as e:
            print(f"❌ Bot Polling Error: {e}")
            time.sleep(5) # ኤረር ቢመጣም ከ5 ሰከንድ በኋላ ቦቱን እንደገና እንዲያነሳው

if __name__ == "__main__":
    bot_thread = threading.Thread(target=run_bot, daemon=True)
    bot_thread.start()
    port = int(os.environ.get("PORT", 10000))
    app.run(host='0.0.0.0', port=port)
